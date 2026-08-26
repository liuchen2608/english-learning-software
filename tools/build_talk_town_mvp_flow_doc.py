from __future__ import annotations

import math
from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "Talk-Town-MVP流程图与任务分工-会议版-v1.0.docx"
ASSET_DIR = ROOT / ".docx_assets_talk_town_mvp"
USER_FLOW_PNG = ASSET_DIR / "talk-town-mvp-user-flow.png"
TIMELINE_PNG = ASSET_DIR / "talk-town-mvp-12-week-timeline.png"

FONT_CJK = "/System/Library/Fonts/STHeiti Light.ttc"
FONT_LATIN = "/System/Library/Fonts/Supplemental/Arial.ttf"

INK = "17342F"
CORAL = "F27052"
MINT = "BDE1CD"
YELLOW = "F2C45E"
PAPER = "F6F3EB"
CARD = "FFFDF8"
LINE = "D8DCD7"
MUTED = "667670"
WHITE = "FFFFFF"
SOFT_GREEN = "EAF4EF"
SOFT_CORAL = "FDEAE5"
SOFT_YELLOW = "FFF4D4"
SOFT_BLUE = "E9F0F5"


def rgb(hex_color: str) -> RGBColor:
    return RGBColor.from_string(hex_color)


def pil_font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont:
    # STHeiti renders Simplified Chinese reliably in the local macOS/LibreOffice stack.
    return ImageFont.truetype(FONT_CJK, size=size, index=1)


def draw_centered_text(
    draw: ImageDraw.ImageDraw,
    box: tuple[int, int, int, int],
    lines: list[str],
    sizes: list[int],
    colors: list[str],
    gaps: int = 8,
) -> None:
    fonts = [pil_font(size) for size in sizes]
    heights = []
    for line, font in zip(lines, fonts):
        bbox = draw.textbbox((0, 0), line, font=font)
        heights.append(bbox[3] - bbox[1])
    total = sum(heights) + gaps * (len(lines) - 1)
    y = box[1] + (box[3] - box[1] - total) / 2
    for line, font, color, height in zip(lines, fonts, colors, heights):
        bbox = draw.textbbox((0, 0), line, font=font)
        width = bbox[2] - bbox[0]
        x = box[0] + (box[2] - box[0] - width) / 2
        draw.text((x, y), line, font=font, fill=f"#{color}")
        y += height + gaps


def draw_arrow(draw: ImageDraw.ImageDraw, start: tuple[int, int], end: tuple[int, int], color: str, width: int = 6) -> None:
    draw.line([start, end], fill=f"#{color}", width=width)
    angle = math.atan2(end[1] - start[1], end[0] - start[0])
    head = 18
    for delta in (2.55, -2.55):
        p = (
            end[0] + head * math.cos(angle + delta),
            end[1] + head * math.sin(angle + delta),
        )
        draw.line([end, p], fill=f"#{color}", width=width)


def build_user_flow() -> None:
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    img = Image.new("RGB", (2400, 760), f"#{CARD}")
    draw = ImageDraw.Draw(img)
    draw.rounded_rectangle((24, 24, 2376, 736), radius=42, fill=f"#{CARD}", outline=f"#{LINE}", width=4)

    nodes = [
        ("01", "输入旅行目标", "中文描述目的地与需求"),
        ("02", "AI 结构化确认", "低置信时继续追问"),
        ("03", "生成学习路线", "只推荐已发布场景"),
        ("04", "咖啡店任务初始化", "固定任务包与版本"),
        ("05", "A / B / C 能力验证", "基线 → 引导 → 迁移"),
        ("06", "结果与学习记录", "比较四项必要意图"),
    ]
    colors = [MINT, SOFT_BLUE, SOFT_YELLOW, SOFT_CORAL, MINT, SOFT_BLUE]
    x0, y0, width, height, gap = 78, 110, 330, 248, 54

    for index, ((num, title, sub), fill) in enumerate(zip(nodes, colors)):
        x = x0 + index * (width + gap)
        draw.rounded_rectangle((x, y0, x + width, y0 + height), radius=30, fill=f"#{fill}", outline=f"#{INK}", width=4)
        draw.ellipse((x + 22, y0 + 22, x + 84, y0 + 84), fill=f"#{INK}")
        nfont = pil_font(24)
        nb = draw.textbbox((0, 0), num, font=nfont)
        draw.text((x + 53 - (nb[2] - nb[0]) / 2, y0 + 53 - (nb[3] - nb[1]) / 2 - 2), num, font=nfont, fill=f"#{YELLOW}")
        draw_centered_text(draw, (x + 18, y0 + 85, x + width - 18, y0 + height - 20), [title, sub], [28, 20], [INK, MUTED], gaps=16)
        if index < len(nodes) - 1:
            draw_arrow(draw, (x + width + 8, y0 + height // 2), (x + width + gap - 8, y0 + height // 2), CORAL, 7)

    bar_y = 455
    draw.rounded_rectangle((190, bar_y, 2210, bar_y + 145), radius=28, fill=f"#{INK}")
    draw_centered_text(
        draw,
        (215, bar_y + 10, 2185, bar_y + 135),
        ["知识与 RAG 支撑", "已审核知识 → 混合检索 → 条件 Rerank → 生成校验 → Trace / 成本"],
        [25, 22],
        [YELLOW, WHITE],
        gaps=13,
    )
    draw_arrow(draw, (1200, bar_y - 5), (1200, y0 + height + 8), CORAL, 7)

    note_font = pil_font(20)
    note = "控制原则：AI 负责理解与反馈；确定性状态机负责最终推进。"
    bbox = draw.textbbox((0, 0), note, font=note_font)
    draw.text(((2400 - (bbox[2] - bbox[0])) / 2, 648), note, font=note_font, fill=f"#{MUTED}")
    img.save(USER_FLOW_PNG, quality=95)


def build_timeline() -> None:
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    img = Image.new("RGB", (2400, 770), f"#{CARD}")
    draw = ImageDraw.Draw(img)
    draw.rounded_rectangle((24, 24, 2376, 746), radius=42, fill=f"#{CARD}", outline=f"#{LINE}", width=4)

    left = 430
    top = 120
    col_w = 145
    row_h = 92
    label_font = pil_font(23)
    week_font = pil_font(20)
    small_font = pil_font(18)

    for week in range(1, 13):
        x = left + (week - 1) * col_w
        fill = PAPER if week % 2 else CARD
        draw.rectangle((x, top - 54, x + col_w, top + row_h * 5), fill=f"#{fill}")
        draw.line((x, top - 54, x, top + row_h * 5), fill=f"#{LINE}", width=2)
        text = f"W{week}"
        bbox = draw.textbbox((0, 0), text, font=week_font)
        draw.text((x + (col_w - (bbox[2] - bbox[0])) / 2, top - 43), text, font=week_font, fill=f"#{MUTED}")
    draw.line((left + 12 * col_w, top - 54, left + 12 * col_w, top + row_h * 5), fill=f"#{LINE}", width=2)

    phases = [
        ("方案冻结", 1, 2, CORAL, "Spike / 原型 / Schema"),
        ("知识摄取与后台", 3, 5, INK, "上传 / 解析 / 审核 / 发布"),
        ("学习端与 AI/RAG", 6, 8, "3E8B73", "状态机 / 检索 / 生成 / 结果"),
        ("联调与离线评测", 9, 10, "D49A25", "金标 / 性能 / 成本 / 安全"),
        ("用户验证", 11, 12, "5F6FB3", "5 人可用性 + 30 人验证"),
    ]

    for row, (name, start, end, color, output) in enumerate(phases):
        y = top + row * row_h
        draw.line((55, y + row_h, 2315, y + row_h), fill=f"#{LINE}", width=2)
        draw.text((65, y + 16), name, font=label_font, fill=f"#{INK}")
        draw.text((65, y + 49), output, font=small_font, fill=f"#{MUTED}")
        x1 = left + (start - 1) * col_w + 10
        x2 = left + end * col_w - 10
        draw.rounded_rectangle((x1, y + 16, x2, y + 70), radius=20, fill=f"#{color}")
        duration = f"{end - start + 1} 周"
        bbox = draw.textbbox((0, 0), duration, font=label_font)
        draw.text(((x1 + x2 - (bbox[2] - bbox[0])) / 2, y + 29), duration, font=label_font, fill=f"#{WHITE}")
        gate_x = x2
        gate_y = y + 43
        draw.polygon([(gate_x, gate_y - 15), (gate_x + 15, gate_y), (gate_x, gate_y + 15), (gate_x - 15, gate_y)], fill=f"#{YELLOW}", outline=f"#{INK}")

    footer = "各阶段可以准备性并行，但进入下一闸门前必须完成本阶段验收。"
    bbox = draw.textbbox((0, 0), footer, font=small_font)
    draw.text(((2400 - (bbox[2] - bbox[0])) / 2, 684), footer, font=small_font, fill=f"#{MUTED}")
    img.save(TIMELINE_PNG, quality=95)


def set_run_font(run, size: float | None = None, color: str | None = None, bold: bool | None = None, italic: bool | None = None) -> None:
    run.font.name = "Arial Unicode MS"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Arial Unicode MS")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Arial Unicode MS")
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Arial Unicode MS")
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = rgb(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_margins(cell, top=90, start=120, bottom=90, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color=LINE, size=6) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), str(size))
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), color)


def set_cell_fill(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_table_geometry(table, widths_in: list[float], indent_dxa: int = 120) -> None:
    widths_dxa = [round(width * 1440) for width in widths_in]
    total = sum(widths_dxa)
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    layout = tbl_pr.first_child_found_in("w:tblLayout")
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            width = widths_dxa[index]
            cell.width = Inches(width / 1440)
            tc_w = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                cell._tc.get_or_add_tcPr().append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def mark_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_cell_text(cell, text: str, size=8.7, color=INK, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = align
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.08
    run = paragraph.add_run(text)
    set_run_font(run, size=size, color=color, bold=bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(doc, headers: list[str], rows: list[list[str]], widths: list[float], font_size=8.7):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    set_table_borders(table)
    header = table.rows[0]
    mark_repeat_table_header(header)
    for index, text in enumerate(headers):
        set_cell_fill(header.cells[index], INK)
        set_cell_text(header.cells[index], text, size=8.5, color=WHITE, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    for row_index, row_data in enumerate(rows):
        row = table.add_row()
        for col_index, text in enumerate(row_data):
            if row_index % 2 == 1:
                set_cell_fill(row.cells[col_index], "F7F8F5")
            align = WD_ALIGN_PARAGRAPH.CENTER if col_index in (0, len(headers) - 1) else WD_ALIGN_PARAGRAPH.LEFT
            set_cell_text(row.cells[col_index], text, size=font_size, align=align)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_before = Pt(0)
    spacer.paragraph_format.space_after = Pt(0)
    spacer.paragraph_format.line_spacing = 0.2
    set_run_font(spacer.add_run("\u200b"), size=1, color=WHITE)
    return table


def add_page_number(paragraph) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for element in (begin, instr, separate, text, end):
        run._r.append(element)
    set_run_font(run, size=8, color=MUTED)


def add_kicker(doc, text: str, page_break_before: bool = False):
    p = doc.add_paragraph()
    p.paragraph_format.page_break_before = page_break_before
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text.upper())
    set_run_font(run, size=8.5, color=CORAL, bold=True)
    return p


def add_title(doc, title: str, subtitle: str | None = None) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(title)
    set_run_font(run, size=23, color=INK, bold=True)
    if subtitle:
        sp = doc.add_paragraph()
        sp.paragraph_format.space_before = Pt(0)
        sp.paragraph_format.space_after = Pt(6)
        sr = sp.add_run(subtitle)
        set_run_font(sr, size=10, color=MUTED)


def add_heading(doc, text: str, level: int = 1) -> None:
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    set_run_font(r, size=13.5 if level == 1 else 11, color=INK if level == 1 else MUTED, bold=True)


def add_body(doc, text: str, bold_lead: str | None = None, after=5) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.12
    if bold_lead and text.startswith(bold_lead):
        first = p.add_run(bold_lead)
        set_run_font(first, size=9.5, color=INK, bold=True)
        rest = p.add_run(text[len(bold_lead):])
        set_run_font(rest, size=9.5, color=INK)
    else:
        run = p.add_run(text)
        set_run_font(run, size=9.5, color=INK)


def add_callout(doc, label: str, text: str, fill=SOFT_YELLOW) -> None:
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [10.0])
    set_table_borders(table, color=fill, size=4)
    cell = table.cell(0, 0)
    set_cell_fill(cell, fill)
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.08
    r1 = p.add_run(f"{label}  ")
    set_run_font(r1, size=9, color=INK, bold=True)
    r2 = p.add_run(text)
    set_run_font(r2, size=9, color=INK)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_before = Pt(0)
    spacer.paragraph_format.space_after = Pt(0)
    spacer.paragraph_format.line_spacing = 0.2
    set_run_font(spacer.add_run("\u200b"), size=1, color=WHITE)


def add_three_scope_cards(doc) -> None:
    table = doc.add_table(rows=1, cols=3)
    set_table_geometry(table, [3.28, 3.28, 3.44])
    set_table_borders(table, color=LINE, size=5)
    cards = [
        ("P0 · 学习端", "目标输入与确认 · 学习路线\n基线 A / 引导 B / 迁移 C · 结果记录", SOFT_GREEN),
        ("P0 · 知识与 AI", "解析 / OCR / 表格 · 审核发布\n混合检索 / Rerank · 生成评测 / Trace", SOFT_BLUE),
        ("本期不做", "语音 / 发音 · 多场景完整训练\n跨设备账户 / 付费 · 长期 Memory / Agent", SOFT_CORAL),
    ]
    row_pr = table.rows[0]._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    row_pr.append(cant_split)
    for cell, (title, body, fill) in zip(table.rows[0].cells, cards):
        set_cell_fill(cell, fill)
        set_cell_margins(cell, top=70, start=110, bottom=70, end=110)
        cell.text = ""
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(title)
        set_run_font(r, size=9.4, color=INK, bold=True)
        bp = cell.add_paragraph()
        bp.paragraph_format.space_before = Pt(0)
        bp.paragraph_format.space_after = Pt(0)
        bp.paragraph_format.line_spacing = 1.0
        br = bp.add_run(body)
        set_run_font(br, size=7.7, color=INK)


def create_document() -> None:
    build_user_flow()
    build_timeline()

    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Inches(11)
    section.page_height = Inches(8.5)
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)
    section.header_distance = Inches(0.25)
    section.footer_distance = Inches(0.25)

    normal = doc.styles["Normal"]
    normal.font.name = "Arial Unicode MS"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial Unicode MS")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial Unicode MS")
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial Unicode MS")
    normal.font.size = Pt(9.5)
    normal.font.color.rgb = rgb(INK)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.12

    for style_name, size, color, before, after in (
        ("Heading 1", 15, INK, 7, 5),
        ("Heading 2", 11.5, MUTED, 5, 3),
        ("Heading 3", 10, MUTED, 4, 2),
    ):
        style = doc.styles[style_name]
        style.font.name = "Arial Unicode MS"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Arial Unicode MS")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial Unicode MS")
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial Unicode MS")
        style.font.size = Pt(size)
        style.font.color.rgb = rgb(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    # Real Word list definitions, with compact meeting-guide geometry.
    for style_name in ("List Bullet", "List Number"):
        style = doc.styles[style_name]
        style.font.name = "Arial Unicode MS"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Arial Unicode MS")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial Unicode MS")
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial Unicode MS")
        style.font.size = Pt(9)
        style.paragraph_format.left_indent = Inches(0.34)
        style.paragraph_format.first_line_indent = Inches(-0.17)
        style.paragraph_format.space_after = Pt(3)
        style.paragraph_format.line_spacing = 1.12

    header_table = section.header.add_table(rows=1, cols=2, width=Inches(10))
    set_table_geometry(header_table, [6.7, 3.3], indent_dxa=0)
    set_table_borders(header_table, color=WHITE, size=0)
    set_cell_text(header_table.cell(0, 0), "TALK TOWN  |  MVP 流程与任务分工", size=7.5, color=MUTED, bold=True)
    set_cell_text(header_table.cell(0, 1), "会议讨论版 · v1.0 · 2026-08-22", size=7.5, color=MUTED, align=WD_ALIGN_PARAGRAPH.RIGHT)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    fr = footer.add_run("内部会议材料  ·  ")
    set_run_font(fr, size=8, color=MUTED)
    add_page_number(footer)

    # Page 1
    add_kicker(doc, "MVP DELIVERY BRIEF")
    add_title(doc, "Talk Town MVP 流程图与任务分工", "完整 AI MVP 基线 · 用于小组评审、分工与阶段验收")
    add_callout(doc, "会议目标", "确认 P0 范围、12 周基线、每项任务 Owner、跨角色依赖和阶段闸门；会后所有任务必须有负责人和截止周。")
    add_heading(doc, "1. MVP 用户学习闭环", 1)
    add_body(doc, "核心验证：一次 AI 引导的咖啡店点餐训练，能否让目标用户在等价新任务中，比训练前完成更多必要交流步骤。")
    pic = doc.add_picture(str(USER_FLOW_PNG), width=Inches(9.55))
    pic._inline.docPr.set("descr", "Talk Town MVP 用户学习闭环流程图")
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.paragraphs[-1].paragraph_format.space_after = Pt(7)
    add_three_scope_cards(doc)

    # Page 2
    add_kicker(doc, "DELIVERY FLOW", page_break_before=True)
    add_title(doc, "12 周交付流程", "阶段可以准备性并行；下一阶段启动前必须通过当前闸门")
    pic = doc.add_picture(str(TIMELINE_PNG), width=Inches(9.88))
    pic._inline.docPr.set("descr", "Talk Town 完整 AI MVP 十二周交付甘特图")
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.paragraphs[-1].paragraph_format.space_after = Pt(6)
    phase_rows = [
        ["1", "方案冻结", "W1–W2", "AI 产品经理", "技术 Spike、模型 Bake-off、页面原型、知识 Schema", "G1 方案签字"],
        ["2", "知识摄取与后台", "W3–W5", "后端 / 知识工程", "上传、解析、审核、切块、版本与发布", "G2 知识链路可用"],
        ["3", "学习端与 AI/RAG", "W6–W8", "前端 + 后端/AI", "学习流程、状态机、检索、生成与结果页", "G3 P0 主链路完成"],
        ["4", "联调与离线评测", "W9–W10", "测试 + AI/数据", "金标集、性能、成本、安全、故障降级与修复", "G4 AI 与系统门槛"],
        ["5", "用户验证", "W11–W12", "产品 + 用户研究", "5 人可用性测试、30 人方向性验证、决策报告", "G5 产品化决策"],
    ]
    add_table(doc, ["#", "阶段", "周期", "主责角色", "关键交付", "退出闸门"], phase_rows, [0.4, 1.45, 0.72, 1.35, 4.15, 1.93], font_size=8.3)

    # Page 3
    add_kicker(doc, "WORK BREAKDOWN", page_break_before=True)
    add_title(doc, "任务分解与责任分配", "会中将“建议 Owner”替换为具体姓名，并确认截止周")
    add_callout(doc, "分工规则", "一个任务只能有一个最终 Owner；协作人可以有多个。涉及范围变化、超过 3 个工作日或预算增加超过 10% 的事项必须重新评审。", fill=SOFT_GREEN)
    task_rows = [
        ["T01", "范围与验收基线", "AI 产品经理\n姓名：____", "全体核心成员", "P0/P1、成功指标、非目标、变更规则完成签字", "W1"],
        ["T02", "目标输入与学习流程原型", "产品设计师\n姓名：____", "产品、前端", "U-01 至 U-09 的主流程、空/错/加载状态和 375px 原型", "W2"],
        ["T03", "A/B/C 场景与评分内容", "英语内容负责人\n姓名：____", "产品、AI", "三套等价任务；四项意图、允许表达、错误与提示规则", "W2"],
        ["T04", "模型网关与结构化输出", "AI 工程师\n姓名：____", "后端、产品", "目标解析、回答判定 Schema；路由、重试、Fallback 与成本记录", "W4"],
        ["T05", "知识摄取与管理后台", "知识/后端工程师\n姓名：____", "内容、测试", "文件解析、OCR、表格、审核、Chunk、版本、发布与回滚", "W5"],
        ["T06", "RAG 检索与评测", "AI/RAG 工程师\n姓名：____", "知识、数据", "Dense + Sparse + RRF + 条件 Rerank；检索调试与金标集", "W7"],
        ["T07", "学习网页端", "前端工程师\n姓名：____", "设计、后端", "目标对话框、路线、A/B/C 训练、结果、异常与进度恢复", "W8"],
        ["T08", "会话、状态机与 API", "后端工程师\n姓名：____", "前端、AI", "匿名会话、幂等 Turn、状态推进、版本固定、记录清除", "W8"],
        ["T09", "埋点、Trace 与成本", "数据/后端\n姓名：____", "产品、AI", "学习/AI/RAG/成本事件可追踪，报表可复算", "W9"],
        ["T10", "质量与发布验收", "测试负责人\n姓名：____", "全体核心成员", "功能、AI、性能、安全、权限、故障注入、回滚用例通过", "W10"],
        ["T11", "可用性与方向性验证", "产品/用户研究\n姓名：____", "数据、内容", "5 人观察测试 + 30 个有效样本 + 继续/迭代/停止建议", "W12"],
    ]
    add_table(doc, ["ID", "工作包", "建议 Owner", "关键协作", "Definition of Done", "截止"], task_rows, [0.48, 1.48, 1.35, 1.34, 4.55, 0.8], font_size=7.45)

    # Page 4
    add_kicker(doc, "QUALITY GATES", page_break_before=True)
    add_title(doc, "验收闸门与会议决策", "任何硬失败项出现时，不得以综合分抵消")
    gate_rows = [
        ["G0 立项", "范围、资源、预算负责人确认", "项目负责人到位；12 周基线获批", "不启动排期"],
        ["G1 方案冻结", "原型、知识 Schema、A/B/C 内容、技术 Spike", "产品、内容、技术共同签字", "退回方案阶段"],
        ["G2 知识链路", "导入、解析、审核、发布、回滚可用", "未审核内容不可上线；来源与版本可追踪", "修复后重验"],
        ["G3 主链路", "目标、路线、A/B/C、结果与状态机贯通", "重复提交不重复推进；异常可恢复", "停止联调"],
        ["G4 用户测试", "金标、性能、成本、安全和发布验收完成", "AI 总分及全部硬门槛通过", "仅内部迭代"],
        ["G5 产品化", "5 人测试完成；30 个有效样本", "学习效果达标且系统质量稳定", "原范围迭代或停止扩展"],
    ]
    add_table(doc, ["闸门", "进入条件", "通过条件", "未通过处理"], gate_rows, [1.25, 3.0, 3.65, 2.1], font_size=8.15)

    add_heading(doc, "核心指标", 1)
    metric_rows = [
        ["AI 上线硬门槛", "总分 ≥ 85；关键槽位准确率 ≥ 95%；结构化输出有效率 ≥ 99%；离线错误推进 = 0；严重安全事件 = 0"],
        ["性能与成本", "单轮端到端 P95 ≤ 8 秒；检索 + 精排 P95 ≤ 800ms；完整任务正常目标 ≤ $0.02"],
        ["建议学习验证线（会中确认）", "平均能力增量 ≥ 1 个步骤；个人改善率 ≥ 60%；引导训练完成率 ≥ 80%"],
        ["发布与可追溯", "任意线上回答可定位到模型、Prompt、Chunk、来源和知识版本；回滚 5 分钟内生效"],
    ]
    add_table(doc, ["类别", "会议确认标准"], metric_rows, [2.35, 7.65], font_size=8.35)

    add_heading(doc, "会议结束前必须完成", 1)
    for text in (
        "为 T01–T11 填入唯一 Owner 姓名和截止周；",
        "确认知识来源与版权审批负责人、知识发布审批人；",
        "确认四套金标集的建设负责人：Parsing、Retrieval、Turn、Safety；",
        "确认 130–160 人日资源基线及现金预算重估责任人；",
        "确定下一次 G1 方案评审时间和必须提交的材料。",
    ):
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(1.5)
        r = p.add_run(text)
        set_run_font(r, size=8.5, color=INK)

    add_body(doc, "基线来源：《Talk Town 完整 AI 产品 PRD v1.0》（2026-08-22）。本文件用于会议执行，不替代详细功能、技术与评测规范。", after=0)

    doc.core_properties.title = "Talk Town MVP 流程图与任务分工"
    doc.core_properties.subject = "完整 AI MVP 会议流程、里程碑、任务分工与验收闸门"
    doc.core_properties.author = "Talk Town 项目组"
    doc.core_properties.keywords = "Talk Town, MVP, PRD, 流程图, 任务分工"
    doc.save(OUTPUT)


if __name__ == "__main__":
    create_document()
    print(OUTPUT)
